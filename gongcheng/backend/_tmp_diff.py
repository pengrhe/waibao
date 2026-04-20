import sys, io
from docx import Document

file1 = r'd:\wp\waibao\gongcheng\cankao\BA-XA-WHSC-02-001深圳市虚拟都市网络技术有限公司虚拟都市网吧.docx'
file2 = r'd:\wp\waibao\gongcheng\文体\BA-SY-WHSC-02-471深圳市奇妙之城游乐管理有限公司.docx'

doc1 = Document(file1)
doc2 = Document(file2)

lines = []

lines.append('='*70)
lines.append('生成文档 - 所有非空段落')
lines.append('='*70)
for i, p in enumerate(doc1.paragraphs):
    text = p.text.strip()
    if text:
        lines.append(f'  P{i}: {text[:120]}')

lines.append('')
lines.append('='*70)
lines.append('参考文档 - 所有非空段落')
lines.append('='*70)
for i, p in enumerate(doc2.paragraphs):
    text = p.text.strip()
    if text:
        lines.append(f'  P{i}: {text[:120]}')

lines.append('')
lines.append('='*70)
lines.append('表格内容逐一对比（仅显示差异）')
lines.append('='*70)
for ti in range(min(len(doc1.tables), len(doc2.tables))):
    t1 = doc1.tables[ti]
    t2 = doc2.tables[ti]
    r1 = len(t1.rows)
    r2 = len(t2.rows)
    c1 = len(t1.columns)
    c2 = len(t2.columns)
    if r1 != r2 or c1 != c2:
        lines.append(f'\nTable {ti}: 尺寸不同 - 生成={r1}x{c1}, 参考={r2}x{c2}')
        lines.append(f'  生成文档:')
        for ri, row in enumerate(t1.rows):
            cells = [cell.text.strip()[:40] for cell in row.cells]
            lines.append(f'    Row {ri}: {cells}')
        lines.append(f'  参考文档:')
        for ri, row in enumerate(t2.rows):
            cells = [cell.text.strip()[:40] for cell in row.cells]
            lines.append(f'    Row {ri}: {cells}')
    else:
        diffs = []
        for ri in range(r1):
            for ci in range(c1):
                v1 = t1.rows[ri].cells[ci].text.strip()
                v2 = t2.rows[ri].cells[ci].text.strip()
                if v1 != v2:
                    diffs.append((ri, ci, v1[:40], v2[:40]))
        if diffs:
            lines.append(f'\nTable {ti} ({r1}x{c1}): {len(diffs)} 个单元格不同')
            for ri, ci, v1, v2 in diffs[:10]:
                lines.append(f'  [{ri},{ci}] 生成: "{v1}" vs 参考: "{v2}"')
            if len(diffs) > 10:
                lines.append(f'  ... 还有 {len(diffs)-10} 处差异')

with open(r'd:\wp\waibao\gongcheng\backend\diff_output.txt', 'w', encoding='utf-8') as f:
    f.write('\n'.join(lines))
print('Done')
