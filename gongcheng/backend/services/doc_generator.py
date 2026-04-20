# -*- coding: utf-8 -*-
"""
Word文档生成引擎：基于模板生成"一企一档"核查验收销号记录
- 填充封面字段（企业名称、地址、核销人员、核销日期）
- 插入企业门面照片和廉政监督卡照片
- 动态生成隐患整改记录表行并插入隐患/整改照片
"""
import os
import copy
from pathlib import Path
from typing import Optional

from docx import Document
from docx.shared import Cm, Pt, Emu, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from config import TEMPLATE_DIR, GENERATED_DIR

FONT_NAME = "仿宋_GB2312"
FONT_SIZE = 16
LINE_SPACING_PT = 20


def _set_cell_vertical_align(cell, align="center"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    existing = tcPr.find(qn('w:vAlign'))
    if existing is not None:
        tcPr.remove(existing)
    vAlign = OxmlElement('w:vAlign')
    vAlign.set(qn('w:val'), align)
    tcPr.append(vAlign)


def _set_cell_text(cell, text: str, font_name: str = None, font_size: int = None,
                   bold: bool = False, align=None):
    """Set cell text with formatting matching the reference document."""
    if font_name is None:
        font_name = FONT_NAME
    if font_size is None:
        font_size = FONT_SIZE

    p = cell.paragraphs[0]
    for extra_p in cell.paragraphs[1:]:
        extra_p._element.getparent().remove(extra_p._element)
    for r_elem in list(p._element.findall(qn('w:r'))):
        p._element.remove(r_elem)

    run = p.add_run(text)
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = Pt(font_size)
    run.font.bold = bold

    pf = p.paragraph_format
    pf.line_spacing = Pt(LINE_SPACING_PT)
    if align is not None:
        pf.alignment = align

    _set_cell_vertical_align(cell, "center")


def _clear_cell(cell):
    """Clear all content in a cell while preserving the cell structure."""
    for p in cell.paragraphs:
        for run in p.runs:
            run.text = ""
        p.text = ""


def _fix_exif_orientation(image_path: str):
    """Apply EXIF orientation and return a corrected temp file path (or original if no rotation needed)."""
    from PIL import Image as PILImage, ImageOps
    import tempfile

    with PILImage.open(image_path) as img:
        exif = img.getexif()
        orientation = exif.get(0x0112, 1)
        if orientation == 1:
            return image_path, img.size[0], img.size[1]

        oriented = ImageOps.exif_transpose(img)
        tmp = tempfile.NamedTemporaryFile(suffix=".jpg", delete=False)
        oriented.save(tmp.name, "JPEG", quality=95)
        return tmp.name, oriented.size[0], oriented.size[1]


def _add_image_to_cell(cell, image_path: str, max_width_cm: float = 6.0, max_height_cm: float = 4.5,
                       force_size: bool = False):
    """Insert an image into a table cell, scaling to fit while preserving aspect ratio."""
    if not image_path or not os.path.exists(image_path):
        return
    p = cell.paragraphs[0]
    run = p.add_run()
    try:
        insert_path, img_w, img_h = _fix_exif_orientation(image_path)

        ratio = img_w / img_h
        target_ratio = max_width_cm / max_height_cm

        if ratio >= target_ratio:
            w = max_width_cm
            h = w / ratio
        else:
            h = max_height_cm
            w = h * ratio

        run.add_picture(insert_path, width=Cm(w), height=Cm(h))
    except Exception:
        try:
            run.add_picture(image_path, width=Cm(max_width_cm))
        except Exception:
            run.add_text("[图片加载失败]")


def _clone_row(table, template_row_idx: int):
    """Clone a table row and append it to the table."""
    template_row = table.rows[template_row_idx]
    new_tr = copy.deepcopy(template_row._tr)
    table._tbl.append(new_tr)
    return table.rows[-1]


def _pad_underline_value(value: str, min_trailing: int = 40) -> str:
    """Pad value with trailing spaces so the underline extends visibly."""
    return value + ' ' * min_trailing


def _fill_cover(doc: Document, project_data: dict):
    """Fill the cover page fields, preserving the label(plain) + value(underlined) run structure."""
    field_map = {
        "企业名称": project_data.get("name", ""),
        "详细地址": project_data.get("address", ""),
        "核销人员": project_data.get("inspectors", ""),
        "核销日期": project_data.get("check_date", ""),
    }

    for p in doc.paragraphs:
        text = p.text.strip()

        matched_label = None
        for field_label in field_map:
            if text.startswith(field_label):
                matched_label = field_label
                break

        if matched_label:
            _fill_underline_field(p, field_map[matched_label])


def _fill_underline_field(paragraph, value: str):
    """Update the value run (Run 1) in a label+value paragraph, preserving underline formatting."""
    runs = paragraph.runs
    if len(runs) >= 2:
        runs[1].text = _pad_underline_value(value)
        for r in runs[2:]:
            r.text = ""
    elif runs:
        label_part = runs[0].text.split("：")[0] + "：" if "：" in runs[0].text else ""
        runs[0].text = label_part + value
    else:
        paragraph.text = value


def _fill_photo_tables(doc: Document, facade_photo: Optional[str], card_photo: Optional[str]):
    """Fill the facade and supervision card photo tables."""
    if len(doc.tables) < 2:
        return

    table_facade = doc.tables[0]
    if len(table_facade.rows) > 1 and facade_photo:
        _add_image_to_cell(table_facade.rows[1].cells[0], facade_photo, max_width_cm=13.20, max_height_cm=8.34, force_size=True)

    table_card = doc.tables[1]
    if len(table_card.rows) > 1 and card_photo:
        _add_image_to_cell(table_card.rows[1].cells[0], card_photo, max_width_cm=13.20, max_height_cm=8.34, force_size=True)


def _fill_hazard_table(doc: Document, project_data: dict, hazards: list):
    """Fill the hazard rectification record table."""
    if len(doc.tables) < 3:
        return

    table = doc.tables[2]
    CENTER = WD_ALIGN_PARAGRAPH.CENTER
    LEFT = WD_ALIGN_PARAGRAPH.LEFT
    JUSTIFY = WD_ALIGN_PARAGRAPH.JUSTIFY

    info_row = table.rows[0]
    _set_cell_text(info_row.cells[0],
                   f"场所名称：{project_data.get('name', '')}",
                   align=JUSTIFY)
    check_date = project_data.get("check_date", "")
    _set_cell_text(info_row.cells[4],
                   f"核查时间：{check_date}",
                   align=JUSTIFY)

    template_row_idx = 2
    existing_data_rows = len(table.rows) - 2
    needed = len(hazards)

    while existing_data_rows < needed:
        _clone_row(table, template_row_idx)
        existing_data_rows += 1

    while existing_data_rows > needed and len(table.rows) > 3:
        tr = table.rows[-1]._tr
        table._tbl.remove(tr)
        existing_data_rows -= 1

    for i, hz in enumerate(hazards):
        row = table.rows[i + 2]
        cells = row.cells

        _set_cell_text(cells[0], str(hz.get("seq", i + 1)), align=CENTER)
        _set_cell_text(cells[1], hz.get("hazard_type", ""), align=CENTER)
        _set_cell_text(cells[2], hz.get("description", ""), align=LEFT)

        hazard_photo = hz.get("hazard_photo_path")
        if hazard_photo and os.path.exists(hazard_photo):
            _clear_cell(cells[3])
            _add_image_to_cell(cells[3], hazard_photo, max_width_cm=4.5, max_height_cm=4.0)
            _set_cell_vertical_align(cells[3], "center")

        _set_cell_text(cells[4], hz.get("rectify_status", "已整改"), align=CENTER)

        rectify_photo = hz.get("rectify_photo_path")
        if rectify_photo and os.path.exists(rectify_photo):
            _clear_cell(cells[5])
            _add_image_to_cell(cells[5], rectify_photo, max_width_cm=4.5, max_height_cm=4.0)
            _set_cell_vertical_align(cells[5], "center")

        remark_val = hz.get("remark", "")
        if remark_val == "限期整改":
            remark_val = ""
        _set_cell_text(cells[6], remark_val, align=CENTER)


def generate_document(
    project_data: dict,
    hazards: list,
    facade_photo: Optional[str] = None,
    card_photo: Optional[str] = None,
    output_filename: Optional[str] = None,
) -> str:
    """
    Generate a Word document from the template.

    Args:
        project_data: dict with keys: name, street, address, inspectors, check_date
        hazards: list of dicts with keys: seq, hazard_type, description, risk,
                 hazard_photo_path, rectify_status, rectify_photo_path, remark
        facade_photo: path to facade photo
        card_photo: path to supervision card photo
        output_filename: optional output filename

    Returns:
        Path to the generated document
    """
    template_path = TEMPLATE_DIR / "template.docx"
    if not template_path.exists():
        raise FileNotFoundError(f"模板文件不存在: {template_path}")

    doc = Document(str(template_path))

    _fill_cover(doc, project_data)
    _fill_photo_tables(doc, facade_photo, card_photo)
    _fill_hazard_table(doc, project_data, hazards)

    if not output_filename:
        street = project_data.get("street", "")
        name = project_data.get("name", "")
        output_filename = f"{street}--{name}现场核查验收销号记录（一企一档）.docx"

    safe_name = "".join(c for c in output_filename if c not in r'\/:*?"<>|')
    output_path = GENERATED_DIR / safe_name
    os.makedirs(GENERATED_DIR, exist_ok=True)
    doc.save(str(output_path))

    return str(output_path)


if __name__ == "__main__":
    import sys, io as _io
    sys.stdout = _io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

    from excel_parser import parse_excel

    xlsx = r"d:\wp\gongcheng\doc\附件：3月份建筑施工专项检查安全隐患问题清单.xlsx"
    photo_dir = r"d:\wp\gongcheng\backend\uploads\hazard_photos"

    projects = parse_excel(xlsx, photo_dir)

    if projects:
        p = projects[0]
        proj_dict = {
            "name": p.name, "street": p.street, "address": p.address,
            "contact": p.contact, "phone": p.phone, "category": p.category,
            "check_date": "2025年3月5日",
        }
        hz_list = []
        for h in p.hazards:
            hz_list.append({
                "seq": h.seq, "hazard_type": h.hazard_type,
                "description": h.description, "risk": h.risk,
                "hazard_photo_path": h.hazard_photo_path,
                "rectify_status": "已整改", "rectify_photo_path": None,
                "remark": h.remark,
            })

        out = generate_document(proj_dict, hz_list)
        print(f"文档已生成: {out}")
        print(f"文件大小: {os.path.getsize(out) / 1024:.0f} KB")
