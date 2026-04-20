# -*- coding: utf-8 -*-
"""
文体项目 Word 报告生成器
基于模板 template_wenti.docx 填充场所信息、隐患记录、照片等。

模板表格结构：
  Table 0: 报告编号
  Table 1: 基本信息（项目名称/委托单位/受检单位/地址/排查人员/日期）
  Table 2: 详细信息（单位名称/地址/负责人/电话/面积/楼层/场所类型/仪器/检测内容/依据）
  Table 3: 场所照片(R1) + 营业执照(R3)
  Table 4: 经营许可证(R1)
  Table 5: 隐患排查表标题
  Table 6: 隐患排查记录（序号/隐患照片/检查项目/隐患描述/整改建议/备注）
  Table 7-24: 电气防火检测附表（保持模板原样）
"""
import os
import re
import copy
import tempfile
from pathlib import Path
from typing import Optional, List

from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from PIL import Image as PILImage, ImageOps

from config import TEMPLATE_DIR, GENERATED_DIR

FONT_NAME = "仿宋_GB2312"
FONT_SIZE = 12


VENUE_TYPE_MAP = {
    "网吧": "网吧",
    "游泳池": "游泳池",
    "游泳": "游泳池",
    "无证泳池": "游泳池",
    "剧本杀": "剧本杀",
    "剧本娱乐": "剧本杀",
    "剧本娱乐（剧本杀）": "剧本杀",
    "密室逃脱": "密室逃脱",
    "密室": "密室逃脱",
    "剧本娱乐（密室逃脱）": "密室逃脱",
    "游艺": "游艺场所",
    "游艺场所": "游艺场所",
    "攀岩": "攀岩馆",
    "攀岩馆": "攀岩馆",
    "星级酒店": "星级酒店",
    "酒店": "星级酒店",
    "A级景区": "A级景区",
    "景区": "A级景区",
    "旅行社": "A级景区",
    "演出场所": "演出场所",
    "演出经营单位": "演出经营单位",
    "演出经营": "演出经营单位",
    "歌舞": "歌舞娱乐场所",
    "歌舞娱乐场所": "歌舞娱乐场所",
    "歌舞娱乐": "歌舞娱乐场所",
}

ALL_VENUE_TYPES = [
    "网吧", "游泳池", "剧本杀", "游艺场所", "攀岩馆", "密室逃脱",
    "星级酒店", "A级景区", "演出场所", "演出经营单位", "歌舞娱乐场所",
]


def _set_cell_vertical_align(cell, align="center"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    existing = tcPr.find(qn('w:vAlign'))
    if existing is not None:
        tcPr.remove(existing)
    vAlign = OxmlElement('w:vAlign')
    vAlign.set(qn('w:val'), align)
    tcPr.append(vAlign)


def _set_cell_text(cell, text: str, font_name=None, font_size=None,
                   bold=False, align=None):
    if font_name is None:
        font_name = FONT_NAME
    if font_size is None:
        font_size = FONT_SIZE

    p = cell.paragraphs[0]
    for extra_p in cell.paragraphs[1:]:
        extra_p._element.getparent().remove(extra_p._element)
    for r_elem in list(p._element.findall(qn('w:r'))):
        p._element.remove(r_elem)

    run = p.add_run(text)
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = Pt(font_size)
    run.font.bold = bold

    pf = p.paragraph_format
    pf.line_spacing = Pt(16)
    if align is not None:
        pf.alignment = align

    _set_cell_vertical_align(cell, "center")


def _clear_cell(cell):
    for p in cell.paragraphs:
        for r_elem in list(p._element.findall(qn('w:r'))):
            p._element.remove(r_elem)


def _clear_table_images(table, data_row_start: int, data_row_end: int = None,
                        image_col: int = 1):
    """清除表格指定行范围内的图片占位内容。"""
    end = data_row_end if data_row_end is not None else len(table.rows)
    for ri in range(data_row_start, min(end, len(table.rows))):
        row = table.rows[ri]
        if image_col < len(row.cells):
            _clear_cell(row.cells[image_col])


def _fix_exif_orientation(image_path: str):
    """修正 EXIF 方向标记，返回 (实际路径, 宽, 高)。"""
    try:
        with PILImage.open(image_path) as img:
            exif = img.getexif()
            if not exif:
                return image_path, img.size[0], img.size[1]
            orientation = exif.get(0x0112, 1)
            if orientation == 1:
                return image_path, img.size[0], img.size[1]
            oriented = ImageOps.exif_transpose(img)
            tmp = tempfile.NamedTemporaryFile(suffix=".jpg", delete=False)
            oriented.save(tmp.name, "JPEG", quality=95)
            return tmp.name, oriented.size[0], oriented.size[1]
    except Exception:
        try:
            with PILImage.open(image_path) as img:
                return image_path, img.size[0], img.size[1]
        except Exception:
            return image_path, 0, 0


def _add_image_to_cell(cell, image_path: str, max_width_cm=6.0, max_height_cm=4.5,
                       force_size=False):
    """插入图片，保持宽高比适配到 max_width_cm × max_height_cm 范围内。"""
    if not image_path or not os.path.exists(image_path):
        return
    p = cell.paragraphs[0]
    run = p.add_run()
    try:
        insert_path, img_w, img_h = _fix_exif_orientation(image_path)
        if img_w == 0 or img_h == 0:
            run.add_picture(insert_path, width=Cm(max_width_cm))
            return
        ratio = img_w / img_h
        target_ratio = max_width_cm / max_height_cm
        if ratio >= target_ratio:
            w = max_width_cm
            h = w / ratio
        else:
            h = max_height_cm
            w = h * ratio
        run.add_picture(insert_path, width=Cm(w), height=Cm(h))
    except Exception:
        try:
            run.add_picture(image_path, width=Cm(max_width_cm))
        except Exception:
            run.add_text("[图片加载失败]")


def _clone_row(table, template_row_idx: int):
    template_row = table.rows[template_row_idx]
    new_tr = copy.deepcopy(template_row._tr)
    table._tbl.append(new_tr)
    return table.rows[-1]


def _build_venue_type_text(venue_type: str) -> str:
    """根据场所类型生成带勾选的类型文本。"""
    matched = VENUE_TYPE_MAP.get(venue_type, "")

    parts = []
    for vt in ALL_VENUE_TYPES:
        if vt == matched:
            parts.append(f"☑{vt}")
        else:
            parts.append(f"□{vt}")

    line1 = " ".join(parts[:6])
    line2 = " ".join(parts[6:])
    return f"{line1}\n{line2}"


def _fill_report_code(doc: Document, report_code: str):
    table = doc.tables[0]
    _set_cell_text(table.rows[0].cells[1], report_code,
                   font_name="宋体", font_size=14, bold=True)


def _format_date_cn(raw: str) -> str:
    """将日期转为 'XXXX年XX月XX日' 格式。"""
    if not raw:
        return ""
    s = str(raw).strip()
    m = re.match(r'(\d{4})[.\-/年](\d{1,2})[.\-/月](\d{1,2})', s)
    if m:
        return f"{m.group(1)}年{int(m.group(2)):02d}月{int(m.group(3)):02d}日"
    return s


def _fill_basic_info(doc: Document, project_data: dict):
    table = doc.tables[1]
    INFO_FONT = "宋体"
    INFO_SIZE = 14
    _set_cell_text(table.rows[2].cells[1], project_data.get("name", ""),
                   font_name=INFO_FONT, font_size=INFO_SIZE)
    _set_cell_text(table.rows[3].cells[1], project_data.get("address", ""),
                   font_name=INFO_FONT, font_size=INFO_SIZE)
    inspectors = project_data.get("inspectors", "")
    if inspectors:
        _set_cell_text(table.rows[4].cells[1], inspectors,
                       font_name=INFO_FONT, font_size=INFO_SIZE)
    check_date = _format_date_cn(project_data.get("check_date", ""))
    _set_cell_text(table.rows[5].cells[1], check_date,
                   font_name=INFO_FONT, font_size=INFO_SIZE)


def _fill_detailed_info(doc: Document, project_data: dict):
    table = doc.tables[2]
    name = project_data.get("name", "")
    address = project_data.get("address", "")
    contact = project_data.get("contact", "")
    phone = project_data.get("phone", "")
    area = project_data.get("area", "")
    floor_info = project_data.get("floor_info", "")
    venue_type = project_data.get("category", "")

    DETAIL_FONT = "宋体"
    DETAIL_SIZE = 12
    _set_cell_text(table.rows[0].cells[1], name, font_name=DETAIL_FONT, font_size=DETAIL_SIZE)
    _set_cell_text(table.rows[1].cells[1], address, font_name=DETAIL_FONT, font_size=DETAIL_SIZE)
    _set_cell_text(table.rows[3].cells[1], contact, font_name=DETAIL_FONT, font_size=DETAIL_SIZE)
    _set_cell_text(table.rows[3].cells[3], phone, font_name=DETAIL_FONT, font_size=DETAIL_SIZE)
    _set_cell_text(table.rows[4].cells[1], area, font_name=DETAIL_FONT, font_size=DETAIL_SIZE)
    _set_cell_text(table.rows[4].cells[3], floor_info, font_name=DETAIL_FONT, font_size=DETAIL_SIZE)

    type_text = _build_venue_type_text(venue_type)
    _set_cell_text(table.rows[5].cells[1], type_text,
                   font_name=DETAIL_FONT, font_size=DETAIL_SIZE)


def _fill_photos(doc: Document, photos: dict):
    """
    photos dict keys: facade, license, permit, record_sheet
    """
    table3 = doc.tables[3]

    facade_path = photos.get("facade")
    if facade_path:
        _clear_cell(table3.rows[1].cells[1])
        _add_image_to_cell(table3.rows[1].cells[1], facade_path,
                           max_width_cm=15.0, max_height_cm=11.25)

    license_path = photos.get("license")
    if license_path:
        _clear_cell(table3.rows[3].cells[1])
        _add_image_to_cell(table3.rows[3].cells[1], license_path,
                           max_width_cm=15.0, max_height_cm=11.25)

    if len(doc.tables) > 4:
        table4 = doc.tables[4]
        permit_path = photos.get("permit")
        if permit_path and len(table4.rows) > 1:
            _clear_cell(table4.rows[1].cells[0])
            _add_image_to_cell(table4.rows[1].cells[0], permit_path,
                               max_width_cm=15.5, max_height_cm=22.0, force_size=False)

    record_sheet_path = photos.get("record_sheet")
    if record_sheet_path and len(doc.tables) > 5:
        table5 = doc.tables[5]
        if len(table5.rows) > 1:
            _clear_cell(table5.rows[1].cells[0])
            _add_image_to_cell(table5.rows[1].cells[0], record_sheet_path,
                               max_width_cm=15.5, max_height_cm=22.0, force_size=False)


def _fill_hazard_table(doc: Document, hazards: list):
    if len(doc.tables) < 7:
        return

    table = doc.tables[6]
    CENTER = WD_ALIGN_PARAGRAPH.CENTER
    LEFT = WD_ALIGN_PARAGRAPH.LEFT

    template_row_idx = 1
    existing_data_rows = len(table.rows) - 1
    needed = len(hazards)

    while existing_data_rows < needed:
        _clone_row(table, template_row_idx)
        existing_data_rows += 1

    while existing_data_rows > needed and len(table.rows) > 2:
        tr = table.rows[-1]._tr
        table._tbl.remove(tr)
        existing_data_rows -= 1

    HZ_FONT = "宋体"
    HZ_SIZE = 12

    for i, hz in enumerate(hazards):
        row = table.rows[i + 1]
        cells = row.cells

        _set_cell_text(cells[0], str(hz.get("seq", i + 1)), align=CENTER,
                       font_name=HZ_FONT, font_size=HZ_SIZE)

        hazard_photo = hz.get("hazard_photo_path")
        if hazard_photo and os.path.exists(hazard_photo):
            _clear_cell(cells[1])
            _add_image_to_cell(cells[1], hazard_photo,
                               max_width_cm=4.0, max_height_cm=3.2, force_size=True)
        elif not hz.get("hazard_photo_path"):
            _set_cell_text(cells[1], "/", align=CENTER,
                           font_name=HZ_FONT, font_size=HZ_SIZE)

        _set_cell_text(cells[2], hz.get("hazard_type", ""), align=CENTER,
                       font_name=HZ_FONT, font_size=HZ_SIZE)
        _set_cell_text(cells[3], hz.get("description", ""), align=LEFT,
                       font_name=HZ_FONT, font_size=HZ_SIZE)
        _set_cell_text(cells[4], hz.get("suggestion", ""), align=LEFT,
                       font_name=HZ_FONT, font_size=HZ_SIZE)

        reference = hz.get("reference", "")
        _set_cell_text(cells[5], reference, align=LEFT,
                       font_name=HZ_FONT, font_size=HZ_SIZE)


def _fill_infrared_table(doc: Document, records: list):
    """Table 7: 带电设备红外检测 — 动态行。R3=header, R4=template data row."""
    if len(doc.tables) <= 7:
        return
    table = doc.tables[7]
    infrared = [r for r in records if r["detection_type"] == "infrared"]
    if not infrared:
        _clear_table_images(table, 4, image_col=1)
        return

    template_row_idx = 4
    existing_data_rows = len(table.rows) - 4  # R0-R3 are headers
    needed = len(infrared)

    while existing_data_rows < needed:
        _clone_row(table, template_row_idx)
        existing_data_rows += 1
    while existing_data_rows > needed and len(table.rows) > 5:
        table._tbl.remove(table.rows[-1]._tr)
        existing_data_rows -= 1

    CENTER = WD_ALIGN_PARAGRAPH.CENTER
    IR_FONT = "宋体"
    IR_SIZE = 12
    for i, rec in enumerate(infrared):
        row = table.rows[4 + i]
        _set_cell_text(row.cells[0], rec.get("location", ""),
                       font_name=IR_FONT, font_size=IR_SIZE, align=CENTER)

        photo = rec.get("photo_path", "")
        if photo and os.path.exists(photo):
            _clear_cell(row.cells[1])
            _add_image_to_cell(row.cells[1], photo,
                               max_width_cm=4.9, max_height_cm=3.7, force_size=True)

        _set_cell_text(row.cells[2], rec.get("code", ""),
                       font_name=IR_FONT, font_size=IR_SIZE, align=CENTER)

        temp = rec.get("temperature", "")
        if temp and not temp.endswith("℃"):
            temp = f"{temp}℃"
        _set_cell_text(row.cells[3], temp,
                       font_name=IR_FONT, font_size=IR_SIZE, align=CENTER)

        _set_cell_text(row.cells[5], rec.get("result", ""),
                       font_name=IR_FONT, font_size=IR_SIZE, align=CENTER)


def _fill_ground_resistance_table(doc: Document, records: list):
    """Table 8: 接地电阻测试 — 动态行。R5=header, R6=template data row, R7=备注."""
    if len(doc.tables) <= 8:
        return
    table = doc.tables[8]
    gr_records = [r for r in records if r["detection_type"] == "ground_resistance"]
    if not gr_records:
        _clear_table_images(table, 6, len(table.rows) - 1, image_col=1)
        return

    template_row_idx = 6
    remark_row = table.rows[-1]
    remark_tr = copy.deepcopy(remark_row._tr)
    existing_data_rows = len(table.rows) - 7  # R0-R5=headers, last=备注
    needed = len(gr_records)

    while existing_data_rows < needed:
        _clone_row(table, template_row_idx)
        existing_data_rows += 1
    while existing_data_rows > needed and len(table.rows) > 8:
        tr = table.rows[len(table.rows) - 2]._tr
        table._tbl.remove(tr)
        existing_data_rows -= 1

    CENTER = WD_ALIGN_PARAGRAPH.CENTER
    GR_FONT = "宋体"
    GR_SIZE = 12
    for i, rec in enumerate(gr_records):
        row = table.rows[6 + i]
        _set_cell_text(row.cells[0], rec.get("location", ""),
                       font_name=GR_FONT, font_size=GR_SIZE, align=CENTER)

        photo = rec.get("photo_path", "")
        if photo and os.path.exists(photo):
            _clear_cell(row.cells[1])
            _add_image_to_cell(row.cells[1], photo,
                               max_width_cm=5.7, max_height_cm=4.3, force_size=True)

        _set_cell_text(row.cells[3], rec.get("resistance_value", ""),
                       font_name=GR_FONT, font_size=GR_SIZE, align=CENTER)
        _set_cell_text(row.cells[4], rec.get("result", ""),
                       font_name=GR_FONT, font_size=GR_SIZE, align=CENTER)


def _fill_simple_detection_tables(doc: Document, records: list):
    """Tables 9-15: 单行检测记录（位置/照片/结果/备注）。"""
    type_to_table = {
        "residual_current": 9,
        "insulation": 10,
        "terminal": 11,
        "indoor_wiring": 12,
        "distribution_box": 13,
        "ceiling_wiring": 14,
        "grounding": 15,
    }
    CENTER = WD_ALIGN_PARAGRAPH.CENTER
    SD_FONT = "宋体"
    SD_SIZE = 12

    for det_type, table_idx in type_to_table.items():
        if len(doc.tables) <= table_idx:
            continue
        table = doc.tables[table_idx]
        matching = [r for r in records if r["detection_type"] == det_type]
        data_row = table.rows[1]

        if not matching:
            _clear_table_images(table, 1, image_col=1)
            _set_cell_text(data_row.cells[2], "符合规范要求",
                           font_name=SD_FONT, font_size=SD_SIZE, align=CENTER)
            _set_cell_text(data_row.cells[3], "/",
                           font_name=SD_FONT, font_size=SD_SIZE)
            continue

        rec = matching[0]

        location = rec.get("location", "")
        if location:
            _set_cell_text(data_row.cells[0], location,
                           font_name=SD_FONT, font_size=SD_SIZE, align=CENTER)

        photo = rec.get("photo_path", "")
        if photo and os.path.exists(photo):
            _clear_cell(data_row.cells[1])
            _add_image_to_cell(data_row.cells[1], photo,
                               max_width_cm=6.4, max_height_cm=4.8, force_size=True)

        result = rec.get("result", "") or "符合规范要求"
        _set_cell_text(data_row.cells[2], result,
                       font_name=SD_FONT, font_size=SD_SIZE, align=CENTER)

        remark = rec.get("remark", "") or "/"
        _set_cell_text(data_row.cells[3], remark,
                       font_name=SD_FONT, font_size=SD_SIZE)


def _fill_checklist_tables(doc: Document, checklist_results: list):
    """Tables 17-24: 设置检测结果列宽度并覆盖检测结果。"""
    from services.checklist_defaults import CHECKLIST_ROW_MAP

    result_map = {}
    for cr in checklist_results:
        key = (cr["table_index"], cr["item_seq"])
        result_map[key] = cr["result"]

    CENTER = WD_ALIGN_PARAGRAPH.CENTER

    for table_idx, (data_start, skip_rows) in CHECKLIST_ROW_MAP.items():
        if len(doc.tables) <= table_idx:
            continue
        table = doc.tables[table_idx]
        ncols = len(table.columns)
        result_col = ncols - 1

        for row in table.rows:
            if result_col < len(row.cells):
                tc = row.cells[result_col]._tc
                tcPr = tc.get_or_add_tcPr()
                tcW = tcPr.find(qn('w:tcW'))
                if tcW is None:
                    tcW = OxmlElement('w:tcW')
                    tcPr.append(tcW)
                tcW.set(qn('w:w'), str(int(3 * 567)))
                tcW.set(qn('w:type'), 'dxa')

        if not result_map:
            continue

        seq = 0
        for ri in range(len(table.rows)):
            if ri < data_start or ri in skip_rows:
                continue
            first_text = table.rows[ri].cells[0].text.strip()
            if first_text in ("检测概述", "续上表", "续下表", "序号") or "本次检测" in first_text or "备注" in first_text:
                continue
            if table_idx == 20 and ri in (3, 15):
                continue
            seq += 1
            key = (table_idx, seq)
            if key in result_map:
                _set_cell_text(
                    table.rows[ri].cells[result_col],
                    result_map[key], font_name="宋体", font_size=10.5,
                    align=CENTER,
                )


def _insert_page_break_before_section4(doc: Document):
    """在"四、电气防火检测报告"段落前插入分页符。"""
    for p in doc.paragraphs:
        text = p.text.strip()
        if text.startswith("四") and "电气防火检测报告" in text:
            pPr = p._element.get_or_add_pPr()
            existing = pPr.find(qn('w:pageBreakBefore'))
            if existing is None:
                pb = OxmlElement('w:pageBreakBefore')
                pPr.append(pb)
            break


def generate_wenti_document(
    project_data: dict,
    hazards: list,
    photos: dict = None,
    detection_records: list = None,
    checklist_results: list = None,
    output_filename: Optional[str] = None,
) -> str:
    template_path = TEMPLATE_DIR / "template_wenti.docx"
    if not template_path.exists():
        raise FileNotFoundError(f"文体模板文件不存在: {template_path}")

    doc = Document(str(template_path))

    report_code_raw = project_data.get("report_code", "")
    m = re.match(r'(BA-\w+-\w+-\d+-\d+)', report_code_raw)
    report_code_short = m.group(1) if m else report_code_raw
    _fill_report_code(doc, report_code_short)
    _fill_basic_info(doc, project_data)
    _fill_detailed_info(doc, project_data)

    if photos:
        _fill_photos(doc, photos)

    _fill_hazard_table(doc, hazards)

    _insert_page_break_before_section4(doc)

    det = detection_records or []
    _fill_infrared_table(doc, det)
    _fill_ground_resistance_table(doc, det)
    _fill_simple_detection_tables(doc, det)

    _fill_checklist_tables(doc, checklist_results or [])

    if not output_filename:
        if report_code_raw:
            output_filename = f"{report_code_raw}.docx"
        else:
            name = project_data.get("name", "")
            output_filename = f"{name}.docx"

    safe_name = "".join(c for c in output_filename if c not in r'\/:*?"<>|')
    output_path = GENERATED_DIR / safe_name
    os.makedirs(GENERATED_DIR, exist_ok=True)
    doc.save(str(output_path))

    return str(output_path)
