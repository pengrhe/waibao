# -*- coding: utf-8 -*-
"""
旅业项目（C包）Word 报告生成器
基于模板 template_lvye.docx 填充场所信息、隐患记录、照片等。

模板表格结构：
  Table 0: 报告编号
  Table 1: 基本信息（项目名称/委托单位/受检单位/地址/排查人员/日期）
  Table 2: 详细信息（单位名称/地址/委托单位/负责人/电话/面积/楼层/场所类型/排查内容/依据）
  Table 3: 场所照片(R1) + 经营许可证(R3)
  Table 4: 营业执照(R1)
  Table 5: 现场隐患排查记录表(R1)
  Table 6-8: 隐患1-3（照片/隐患描述/整改建议/备注）
"""
import os
import re
import copy
import tempfile
from pathlib import Path
from typing import Optional, List

from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from PIL import Image as PILImage, ImageOps

from config import TEMPLATE_DIR, GENERATED_DIR

FONT_NAME = "仿宋_GB2312"
FONT_SIZE = 12


def _set_cell_vertical_align(cell, align="center"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    existing = tcPr.find(qn('w:vAlign'))
    if existing is not None:
        tcPr.remove(existing)
    vAlign = OxmlElement('w:vAlign')
    vAlign.set(qn('w:val'), align)
    tcPr.append(vAlign)


def _set_cell_text(cell, text: str, font_name=None, font_size=None,
                   bold=False, align=None):
    if font_name is None:
        font_name = FONT_NAME
    if font_size is None:
        font_size = FONT_SIZE

    p = cell.paragraphs[0]
    for extra_p in cell.paragraphs[1:]:
        extra_p._element.getparent().remove(extra_p._element)
    for r_elem in list(p._element.findall(qn('w:r'))):
        p._element.remove(r_elem)

    run = p.add_run(text)
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = Pt(font_size)
    run.font.bold = bold

    pf = p.paragraph_format
    pf.line_spacing = Pt(16)
    if align is not None:
        pf.alignment = align

    _set_cell_vertical_align(cell, "center")


def _clear_cell(cell):
    for p in cell.paragraphs:
        for r_elem in list(p._element.findall(qn('w:r'))):
            p._element.remove(r_elem)


def _fix_exif_orientation(image_path: str):
    try:
        with PILImage.open(image_path) as img:
            exif = img.getexif()
            if not exif:
                return image_path, img.size[0], img.size[1]
            orientation = exif.get(0x0112, 1)
            if orientation == 1:
                return image_path, img.size[0], img.size[1]
            oriented = ImageOps.exif_transpose(img)
            tmp = tempfile.NamedTemporaryFile(suffix=".jpg", delete=False)
            oriented.save(tmp.name, "JPEG", quality=95)
            return tmp.name, oriented.size[0], oriented.size[1]
    except Exception:
        try:
            with PILImage.open(image_path) as img:
                return image_path, img.size[0], img.size[1]
        except Exception:
            return image_path, 0, 0


def _get_cell_width_cm(cell):
    """从单元格 XML 属性读取实际宽度（cm），无法获取时返回 None。"""
    tc = cell._tc
    tcPr = tc.find(qn('w:tcPr'))
    if tcPr is not None:
        tcW = tcPr.find(qn('w:tcW'))
        if tcW is not None:
            w_val = tcW.get(qn('w:w'))
            w_type = tcW.get(qn('w:type'), 'dxa')
            if w_val and w_type == 'dxa':
                return int(w_val) / 567.0
    return None


def _set_row_height_auto(row):
    """将行高规则从 exact 改为 atLeast，允许行高随内容扩展。"""
    tr = row._tr
    trPr = tr.find(qn('w:trPr'))
    if trPr is not None:
        trHeight = trPr.find(qn('w:trHeight'))
        if trHeight is not None:
            trHeight.set(qn('w:hRule'), 'atLeast')


def _add_image_to_cell(cell, image_path: str, max_width_cm=6.0, max_height_cm=4.5):
    if not image_path or not os.path.exists(image_path):
        return

    cell_w = _get_cell_width_cm(cell)
    if cell_w is not None and cell_w > 1.0:
        effective_w = cell_w - 0.3
        if effective_w < max_width_cm:
            max_width_cm = effective_w

    p = cell.paragraphs[0]
    run = p.add_run()
    try:
        insert_path, img_w, img_h = _fix_exif_orientation(image_path)
        if img_w == 0 or img_h == 0:
            run.add_picture(insert_path, width=Cm(max_width_cm))
            return
        ratio = img_w / img_h
        target_ratio = max_width_cm / max_height_cm
        if ratio >= target_ratio:
            w = max_width_cm
            h = w / ratio
        else:
            h = max_height_cm
            w = h * ratio
        run.add_picture(insert_path, width=Cm(w), height=Cm(h))
    except Exception:
        try:
            run.add_picture(image_path, width=Cm(max_width_cm))
        except Exception:
            run.add_text("[图片加载失败]")


def _clone_table_after(doc: Document, source_table_idx: int) -> int:
    """深拷贝一个表格并插入到源表格之后，返回新表格的索引。"""
    src_tbl = doc.tables[source_table_idx]._tbl
    new_tbl = copy.deepcopy(src_tbl)

    parent = src_tbl.getparent()
    src_tbl.addnext(new_tbl)

    for i, t in enumerate(doc.tables):
        if t._tbl is new_tbl:
            return i
    return len(doc.tables) - 1


def _fill_report_code(doc: Document, report_code: str):
    table = doc.tables[0]
    _set_cell_text(table.rows[0].cells[1], report_code,
                   font_name="宋体", font_size=14, bold=True)


def _format_date_cn(raw: str) -> str:
    if not raw:
        return ""
    s = str(raw).strip()
    m = re.match(r'(\d{4})[.\-/年](\d{1,2})[.\-/月](\d{1,2})', s)
    if m:
        return f"{m.group(1)}年{int(m.group(2)):02d}月{int(m.group(3)):02d}日"
    return s


def _fill_basic_info(doc: Document, project_data: dict):
    table = doc.tables[1]
    INFO_FONT = "宋体"
    INFO_SIZE = 14
    _set_cell_text(table.rows[2].cells[1], project_data.get("name", ""),
                   font_name=INFO_FONT, font_size=INFO_SIZE)
    _set_cell_text(table.rows[3].cells[1], project_data.get("address", ""),
                   font_name=INFO_FONT, font_size=INFO_SIZE)
    inspectors = project_data.get("inspectors", "")
    if inspectors:
        _set_cell_text(table.rows[4].cells[1], inspectors,
                       font_name=INFO_FONT, font_size=INFO_SIZE)
    check_date = _format_date_cn(project_data.get("check_date", ""))
    _set_cell_text(table.rows[5].cells[1], check_date,
                   font_name=INFO_FONT, font_size=INFO_SIZE)


def _fill_detailed_info(doc: Document, project_data: dict):
    table = doc.tables[2]
    DETAIL_FONT = "宋体"
    DETAIL_SIZE = 12

    name = project_data.get("name", "")
    address = project_data.get("address", "")
    contact = project_data.get("contact", "")
    phone = project_data.get("phone", "")
    area = project_data.get("area", "")
    floor_info = project_data.get("floor_info", "")

    _set_cell_text(table.rows[0].cells[1], name,
                   font_name=DETAIL_FONT, font_size=DETAIL_SIZE)
    _set_cell_text(table.rows[1].cells[1], address,
                   font_name=DETAIL_FONT, font_size=DETAIL_SIZE)
    _set_cell_text(table.rows[3].cells[1], contact,
                   font_name=DETAIL_FONT, font_size=DETAIL_SIZE)
    _set_cell_text(table.rows[3].cells[3], phone,
                   font_name=DETAIL_FONT, font_size=DETAIL_SIZE)
    _set_cell_text(table.rows[4].cells[1], area,
                   font_name=DETAIL_FONT, font_size=DETAIL_SIZE)
    _set_cell_text(table.rows[4].cells[3], floor_info,
                   font_name=DETAIL_FONT, font_size=DETAIL_SIZE)


def _fill_photos(doc: Document, photos: dict):
    """
    photos dict keys: facade, license, permit, record_sheet
    Table 3: facade(R1) + permit(R3)
    Table 4: license(R1)
    Table 5: record_sheet(R1)
    """
    table3 = doc.tables[3]

    facade_path = photos.get("facade")
    if facade_path:
        _clear_cell(table3.rows[1].cells[1])
        _add_image_to_cell(table3.rows[1].cells[1], facade_path,
                           max_width_cm=15.0, max_height_cm=11.25)
        _set_row_height_auto(table3.rows[1])

    permit_path = photos.get("permit")
    if permit_path:
        # 经营许可证：模板已将该行合并为整行单元格 (~16.6cm 宽)，
        # 行高 atLeast 8cm，专为横拍许可证设计。
        permit_cell = table3.rows[3].cells[0]
        _clear_cell(permit_cell)
        _add_image_to_cell(permit_cell, permit_path,
                           max_width_cm=16.0, max_height_cm=10.0)
        _set_row_height_auto(table3.rows[3])

    if len(doc.tables) > 4:
        table4 = doc.tables[4]
        license_path = photos.get("license")
        if license_path and len(table4.rows) > 1:
            _clear_cell(table4.rows[1].cells[0])
            _add_image_to_cell(table4.rows[1].cells[0], license_path,
                               max_width_cm=15.5, max_height_cm=22.0)
            _set_row_height_auto(table4.rows[1])

    record_sheet_path = photos.get("record_sheet")
    if record_sheet_path and len(doc.tables) > 5:
        table5 = doc.tables[5]
        if len(table5.rows) > 1:
            _clear_cell(table5.rows[1].cells[0])
            _add_image_to_cell(table5.rows[1].cells[0], record_sheet_path,
                               max_width_cm=15.5, max_height_cm=22.0)
            _set_row_height_auto(table5.rows[1])


def _fill_hazard_tables(doc: Document, hazards: list):
    """
    填充隐患表格（Tables 6-8）。
    模板自带3个隐患表格，若隐患>3则克隆更多，若<3则清空多余的。
    每个隐患表格结构：
      Row0: 标题 "隐患N：" (gridSpan=3)
      Row1: Col0=照片, Col1="隐患描述：", Col2=描述文本
      Row2: Col0=照片(续), Col1="整改建议：", Col2=建议文本
      Row3: Col0=照片(续), Col1="备注：", Col2=备注文本
    """
    HAZARD_TABLE_START = 6
    template_count = 3
    needed = len(hazards)

    HZ_FONT = "宋体"
    HZ_SIZE = 12
    LEFT = WD_ALIGN_PARAGRAPH.LEFT

    while len(doc.tables) - HAZARD_TABLE_START < needed:
        last_hz_idx = len(doc.tables) - 1
        _clone_table_after(doc, last_hz_idx)

    hz_tables = list(range(HAZARD_TABLE_START, HAZARD_TABLE_START + needed))

    for i, hz in enumerate(hazards):
        table_idx = hz_tables[i]
        table = doc.tables[table_idx]
        seq = hz.get("seq", i + 1)

        title_cell = table.rows[0].cells[0]
        _set_cell_text(title_cell, f"隐患{seq}：",
                       font_name=HZ_FONT, font_size=HZ_SIZE, bold=True)

        hazard_photo = hz.get("hazard_photo_path")
        if hazard_photo and os.path.exists(hazard_photo):
            _clear_cell(table.rows[1].cells[0])
            _add_image_to_cell(table.rows[1].cells[0], hazard_photo,
                               max_width_cm=5.0, max_height_cm=7.0)

        _set_cell_text(table.rows[1].cells[2], hz.get("description", ""),
                       font_name=HZ_FONT, font_size=HZ_SIZE, align=LEFT)

        _set_cell_text(table.rows[2].cells[2], hz.get("suggestion", ""),
                       font_name=HZ_FONT, font_size=HZ_SIZE, align=LEFT)

        remark = hz.get("remark", "")
        _set_cell_text(table.rows[3].cells[2], remark,
                       font_name=HZ_FONT, font_size=HZ_SIZE, align=LEFT)

    for j in range(needed, min(template_count, len(doc.tables) - HAZARD_TABLE_START)):
        table = doc.tables[HAZARD_TABLE_START + j]
        for ri in range(1, len(table.rows)):
            for ci in range(len(table.rows[ri].cells)):
                _clear_cell(table.rows[ri].cells[ci])


def generate_lvye_document(
    project_data: dict,
    hazards: list,
    photos: dict = None,
    output_filename: Optional[str] = None,
) -> str:
    template_path = TEMPLATE_DIR / "template_lvye.docx"
    if not template_path.exists():
        raise FileNotFoundError(f"旅业模板文件不存在: {template_path}")

    doc = Document(str(template_path))

    report_code = project_data.get("report_code", "")
    _fill_report_code(doc, report_code)
    _fill_basic_info(doc, project_data)
    _fill_detailed_info(doc, project_data)

    if photos:
        _fill_photos(doc, photos)

    _fill_hazard_tables(doc, hazards)

    if not output_filename:
        if report_code:
            name = project_data.get("name", "")
            output_filename = f"{report_code}{name}.docx"
        else:
            name = project_data.get("name", "")
            output_filename = f"{name}.docx"

    safe_name = "".join(c for c in output_filename if c not in r'\/:*?"<>|')
    output_path = GENERATED_DIR / safe_name
    os.makedirs(GENERATED_DIR, exist_ok=True)
    doc.save(str(output_path))

    return str(output_path)
