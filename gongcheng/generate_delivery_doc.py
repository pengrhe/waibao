# -*- coding: utf-8 -*-
"""
Generate client delivery document.
Run: python generate_delivery_doc.py
Output: ./系统功能说明书.docx
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

QR_IMAGE = r"C:\Users\admin\Desktop\极速办公帮体验版（4月14日前有效）.png"
OUTPUT = os.path.join(os.path.dirname(os.path.abspath(__file__)), "\u7cfb\u7edf\u529f\u80fd\u8bf4\u660e\u4e66.docx")
ADMIN_URL = "https://me6lyl8xtw9bg7x.shanxiangjiaoyu.com/gongcheng"

LQ = "\u201c"  # left double quotation mark
RQ = "\u201d"  # right double quotation mark
ARROW = "\u2192"


def set_cell_shading(cell, color):
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge, val in kwargs.items():
        el = OxmlElement("w:" + edge)
        el.set(qn("w:val"), val.get("val", "single"))
        el.set(qn("w:sz"), val.get("sz", "4"))
        el.set(qn("w:color"), val.get("color", "CCCCCC"))
        el.set(qn("w:space"), "0")
        tcBorders.append(el)
    tcPr.append(tcBorders)


def add_para(doc, text, fname="微软雅黑", size=11, bold=False,
             color=None, align=None, sb=0, sa=6, ls=1.35):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = fname
    run._element.rPr.rFonts.set(qn("w:eastAsia"), fname)
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    pf = p.paragraph_format
    if align is not None:
        pf.alignment = align
    pf.space_before = Pt(sb)
    pf.space_after = Pt(sa)
    pf.line_spacing = ls
    return p, run


def add_h(doc, text, level=1, color=(55, 48, 163)):
    sizes = {1: 22, 2: 16, 3: 13}
    befores = {1: 24, 2: 18, 3: 12}
    return add_para(doc, text, size=sizes.get(level, 13), bold=True,
                    color=color, sb=befores.get(level, 12), sa=8)


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.clear()
    run = p.add_run(text)
    run.font.name = "\u5fae\u8f6f\u96c5\u9ed1"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "\u5fae\u8f6f\u96c5\u9ed1")
    run.font.size = Pt(10.5)
    pf = p.paragraph_format
    pf.space_before = Pt(2)
    pf.space_after = Pt(2)
    pf.line_spacing = 1.3
    return p


def make_table(doc, title, features, accent="4338CA"):
    if title:
        add_h(doc, title, level=3)

    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    hdr = table.rows[0].cells
    for i, txt in enumerate(["\u529f\u80fd\u6a21\u5757", "\u529f\u80fd\u8bf4\u660e"]):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        run = p.add_run(txt)
        run.font.name = "\u5fae\u8f6f\u96c5\u9ed1"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "\u5fae\u8f6f\u96c5\u9ed1")
        run.font.size = Pt(10)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(hdr[i], accent)

    for module, desc in features:
        row = table.add_row().cells
        for i, txt in enumerate([module, desc]):
            row[i].text = ""
            p = row[i].paragraphs[0]
            run = p.add_run(txt)
            run.font.name = "\u5fae\u8f6f\u96c5\u9ed1"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "\u5fae\u8f6f\u96c5\u9ed1")
            run.font.size = Pt(9.5)
            pf = p.paragraph_format
            pf.space_before = Pt(3)
            pf.space_after = Pt(3)
            pf.line_spacing = 1.25

    widths = [Cm(3.5), Cm(13)]
    for row in table.rows:
        for i, w in enumerate(widths):
            row.cells[i].width = w

    border_style = {"val": "single", "sz": "4", "color": "E2E8F0"}
    for ri, row in enumerate(table.rows):
        for cell in row.cells:
            set_cell_border(cell,
                            top=border_style, bottom=border_style,
                            left=border_style, right=border_style)
            if ri > 0 and ri % 2 == 0:
                set_cell_shading(cell, "F8FAFC")

    return table


def q(text):
    return LQ + text + RQ


def build():
    doc = Document()

    style = doc.styles["Normal"]
    font = style.font
    font.name = "\u5fae\u8f6f\u96c5\u9ed1"
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "\u5fae\u8f6f\u96c5\u9ed1")
    font.size = Pt(10.5)

    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.8)
        section.right_margin = Cm(2.8)

    # ===== COVER =====
    for _ in range(4):
        doc.add_paragraph()

    add_para(doc, "\u5efa\u7b51\u65bd\u5de5\u5b89\u5168\u9690\u60a3\u68c0\u67e5\u81ea\u52a8\u5316\u7cfb\u7edf",
             size=26, bold=True, color=(30, 27, 75),
             align=WD_ALIGN_PARAGRAPH.CENTER, sa=8)

    add_para(doc, "\u529f \u80fd \u8bf4 \u660e \u4e66",
             size=18, color=(99, 102, 241),
             align=WD_ALIGN_PARAGRAPH.CENTER, sa=30)

    add_para(doc, "\u2014\u2014  Excel \u89e3\u6790 \u00b7 \u73b0\u573a\u53d6\u8bc1 \u00b7 Word \u6587\u6863\u81ea\u52a8\u751f\u6210  \u2014\u2014",
             size=11, color=(148, 163, 184),
             align=WD_ALIGN_PARAGRAPH.CENTER, sa=60)

    info_lines = [
        ("\u9879\u76ee\u540d\u79f0", "\u5efa\u7b51\u65bd\u5de5\u5b89\u5168\u9690\u60a3\u68c0\u67e5\u81ea\u52a8\u5316\u7cfb\u7edf"),
        ("\u7cfb\u7edf\u7248\u672c", "V1.0"),
        ("\u4ea4\u4ed8\u65e5\u671f", "2025\u5e744\u6708"),
        ("\u6280\u672f\u652f\u6301", "\u6781\u901f\u529e\u516c\u5e2e"),
    ]
    table = doc.add_table(rows=len(info_lines), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (label, value) in enumerate(info_lines):
        c0, c1 = table.cell(i, 0), table.cell(i, 1)
        c0.width, c1.width = Cm(3), Cm(8)
        for cell, txt, bld in [(c0, label, True), (c1, value, False)]:
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(txt)
            run.font.name = "\u5fae\u8f6f\u96c5\u9ed1"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "\u5fae\u8f6f\u96c5\u9ed1")
            run.font.size = Pt(11)
            run.font.bold = bld
            run.font.color.rgb = RGBColor(71, 85, 105) if bld else RGBColor(15, 23, 42)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)

    doc.add_page_break()

    # ===== TOC =====
    add_h(doc, "\u76ee  \u5f55", level=1, color=(30, 27, 75))
    toc = [
        (True, "\u4e00\u3001\u8bbf\u95ee\u65b9\u5f0f"),
        (True, "\u4e8c\u3001\u7cfb\u7edf\u7b80\u4ecb"),
        (True, "\u4e09\u3001\u7cfb\u7edf\u67b6\u6784"),
        (True, "\u56db\u3001\u540e\u53f0\u7ba1\u7406\u7aef\u529f\u80fd"),
        (False, "    4.1  \u5de5\u4f5c\u53f0\uff08\u6570\u636e\u603b\u89c8\uff09"),
        (False, "    4.2  \u4e0a\u4f20\u9690\u60a3\u6e05\u5355"),
        (False, "    4.3  \u9879\u76ee\u7ba1\u7406"),
        (False, "    4.4  \u9879\u76ee\u8be6\u60c5"),
        (False, "    4.5  \u6587\u6863\u751f\u6210\u4e0e\u4e0b\u8f7d"),
        (True, "\u4e94\u3001\u5fae\u4fe1\u5c0f\u7a0b\u5e8f\u7aef\u529f\u80fd"),
        (False, "    5.1  \u68c0\u67e5\u4efb\u52a1\u5217\u8868"),
        (False, "    5.2  \u9879\u76ee\u8be6\u60c5"),
        (False, "    5.3  \u9690\u60a3\u8be6\u60c5\u4e0e\u6574\u6539"),
        (False, "    5.4  \u73b0\u573a\u62cd\u7167\u53d6\u8bc1"),
        (True, "\u516d\u3001\u6838\u5fc3\u6280\u672f\u4eae\u70b9"),
    ]
    for bld, item in toc:
        add_para(doc, item, size=11 if bld else 10.5, bold=bld,
                 color=(30, 27, 75) if bld else (71, 85, 105), sb=2, sa=2)
    doc.add_page_break()

    # ===== 1. ACCESS =====
    add_h(doc, "\u4e00\u3001\u8bbf\u95ee\u65b9\u5f0f", level=1)

    add_h(doc, "1.1  \u540e\u53f0\u7ba1\u7406\u7aef", level=2)
    add_para(doc, "\u540e\u53f0\u7ba1\u7406\u7aef\u901a\u8fc7\u6d4f\u89c8\u5668\u8bbf\u95ee\uff0c\u5efa\u8bae\u4f7f\u7528 Chrome\u3001Edge \u7b49\u73b0\u4ee3\u6d4f\u89c8\u5668\u3002",
             size=11, sb=6, sa=6)
    add_para(doc, "\u8bbf\u95ee\u5730\u5740\uff1a" + ADMIN_URL, size=11, color=(99, 102, 241), sa=12, bold=True)

    for step in [
        "1. \u5728\u6d4f\u89c8\u5668\u5730\u5740\u680f\u8f93\u5165\u4e0a\u8ff0\u5730\u5740\u5e76\u56de\u8f66",
        "2. \u8fdb\u5165\u7cfb\u7edf\u540e\uff0c\u5de6\u4fa7\u83dc\u5355\u53ef\u5207\u6362\u5de5\u4f5c\u53f0\u3001\u4e0a\u4f20\u6e05\u5355\u3001\u9879\u76ee\u7ba1\u7406\u4e09\u4e2a\u529f\u80fd\u9875\u9762",
        "3. \u4e0a\u4f20 Excel \u6e05\u5355\u540e\uff0c\u7cfb\u7edf\u81ea\u52a8\u89e3\u6790\u5e76\u5728\u9879\u76ee\u7ba1\u7406\u4e2d\u5c55\u793a",
        "4. \u9879\u76ee\u5168\u90e8\u5b8c\u6210\u540e\uff0c\u53ef\u6279\u91cf\u751f\u6210 Word \u68c0\u67e5\u6863\u6848\u5e76\u4e0b\u8f7d",
    ]:
        add_bullet(doc, step)

    add_h(doc, "1.2  \u5fae\u4fe1\u5c0f\u7a0b\u5e8f\u7aef", level=2)
    add_para(doc, "\u5fae\u4fe1\u5c0f\u7a0b\u5e8f\u7aef\u9762\u5411\u4e00\u7ebf\u68c0\u67e5\u4eba\u5458\uff0c\u5fae\u4fe1\u626b\u63cf\u4e0b\u65b9\u4e8c\u7ef4\u7801\u5373\u53ef\u4f53\u9a8c\u3002",
             size=11, sb=6, sa=8)

    if os.path.exists(QR_IMAGE):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(QR_IMAGE, width=Cm(5))

        add_para(doc, "\u5fae\u4fe1\u626b\u7801\u4f53\u9a8c\u5c0f\u7a0b\u5e8f",
                 size=10, color=(148, 163, 184),
                 align=WD_ALIGN_PARAGRAPH.CENTER, sb=4, sa=4)
        add_para(doc, "\uff08\u4f53\u9a8c\u7248\u4e8c\u7ef4\u7801\uff0c\u6709\u6548\u671f\u81f3 2025 \u5e74 4 \u6708 14 \u65e5\uff09",
                 size=9, color=(148, 163, 184),
                 align=WD_ALIGN_PARAGRAPH.CENTER, sa=12)

    for step in [
        "1. \u6253\u5f00\u5fae\u4fe1\uff0c\u626b\u63cf\u4e0a\u65b9\u4e8c\u7ef4\u7801\u8fdb\u5165\u5c0f\u7a0b\u5e8f",
        "2. \u9996\u9875\u5c55\u793a\u6240\u6709\u68c0\u67e5\u4efb\u52a1\uff0c\u6309\u8857\u9053\u5206\u7ec4\uff0c\u652f\u6301\u641c\u7d22",
        "3. \u70b9\u51fb\u9879\u76ee\u8fdb\u5165\u8be6\u60c5\uff0c\u67e5\u770b\u9690\u60a3\u5217\u8868\u548c\u7167\u7247",
        "4. \u70b9\u51fb\u7167\u7247\u533a\u57df\u53ef\u8df3\u8f6c\u62cd\u7167\u9875\u9762\uff0c\u62cd\u6444\u95e8\u9762\u7167\u3001\u76d1\u7763\u5361\u548c\u6574\u6539\u7167\u7247",
        "5. \u62cd\u6444\u7684\u7167\u7247\u81ea\u52a8\u53e0\u52a0\u6c34\u5370\u5e76\u4e0a\u4f20\u81f3\u670d\u52a1\u5668",
        "6. \u6240\u6709\u7167\u7247\u4e0a\u4f20\u5b8c\u6210\u540e\uff0c\u9879\u76ee\u81ea\u52a8\u6807\u8bb0\u4e3a" + q("\u5df2\u5b8c\u6210"),
    ]:
        add_bullet(doc, step)

    # ===== 2. INTRO =====
    add_h(doc, "\u4e8c\u3001\u7cfb\u7edf\u7b80\u4ecb", level=1)
    add_para(doc,
             "\u5efa\u7b51\u65bd\u5de5\u5b89\u5168\u9690\u60a3\u68c0\u67e5\u81ea\u52a8\u5316\u7cfb\u7edf"
             "\u662f\u4e00\u5957\u9762\u5411\u653f\u5e9c\u5b89\u5168\u76d1\u7ba1\u90e8\u95e8\u7684\u6570\u5b57\u5316\u89e3\u51b3\u65b9\u6848\uff0c"
             "\u65e8\u5728\u5c06\u4f20\u7edf\u7684\u5efa\u7b51\u65bd\u5de5\u5b89\u5168\u9690\u60a3\u6392\u67e5\u5de5\u4f5c\u4ece"
             + LQ + "\u7eb8\u8d28\u53f0\u8d26 + \u4eba\u5de5\u6574\u7406" + RQ + "\u5347\u7ea7\u4e3a"
             + LQ + "\u81ea\u52a8\u89e3\u6790 + \u79fb\u52a8\u53d6\u8bc1 + \u4e00\u952e\u751f\u6210" + RQ
             + "\u7684\u5168\u6d41\u7a0b\u81ea\u52a8\u5316\u6a21\u5f0f\u3002",
             size=11, sb=6, sa=8)
    add_para(doc,
             "\u7cfb\u7edf\u5305\u542b\u540e\u53f0\u7ba1\u7406\u7aef\uff08Web\uff09\u548c\u5fae\u4fe1\u5c0f\u7a0b\u5e8f\u7aef\u4e24\u4e2a\u4f7f\u7528\u5165\u53e3\uff0c"
             "\u8986\u76d6\u4ece\u9690\u60a3\u6e05\u5355\u5bfc\u5165\u3001\u73b0\u573a\u7167\u7247\u91c7\u96c6\u3001\u6574\u6539\u8ddf\u8e2a\u5230"
             "\u68c0\u67e5\u6863\u6848\uff08Word\uff09\u81ea\u52a8\u751f\u6210\u7684\u5b8c\u6574\u4e1a\u52a1\u95ed\u73af\u3002",
             size=11, sa=8)

    for h in [
        "\u4e00\u952e\u5bfc\u5165\uff1a\u4e0a\u4f20 Excel \u9690\u60a3\u6e05\u5355\uff0c\u81ea\u52a8\u89e3\u6790\u9879\u76ee\u4fe1\u606f\u3001\u9690\u60a3\u8bb0\u5f55\u53ca\u5d4c\u5165\u7167\u7247",
        "\u79fb\u52a8\u53d6\u8bc1\uff1a\u5fae\u4fe1\u5c0f\u7a0b\u5e8f\u73b0\u573a\u62cd\u7167\uff0c\u81ea\u52a8\u53e0\u52a0\u6c34\u5370\uff08\u673a\u6784\u540d\u79f0\u3001\u65f6\u95f4\u3001\u5730\u70b9\uff09",
        "\u72b6\u6001\u8ddf\u8e2a\uff1a\u5b9e\u65f6\u5c55\u793a\u5404\u9879\u76ee" + LQ + "\u5f85\u5904\u7406 " + ARROW + " \u8fdb\u884c\u4e2d " + ARROW + " \u5df2\u5b8c\u6210" + RQ + "\u72b6\u6001\u4e0e\u7167\u7247\u8fdb\u5ea6",
        "\u4e00\u952e\u751f\u6210\uff1a\u81ea\u52a8\u586b\u5145 Word \u6a21\u677f\uff0c\u751f\u6210" + q("\u4e00\u4f01\u4e00\u6863") + "\u68c0\u67e5\u6863\u6848\uff0c\u5373\u523b\u4e0b\u8f7d",
        "\u6279\u91cf\u64cd\u4f5c\uff1a\u652f\u6301\u6279\u91cf\u751f\u6210\u6587\u6863\uff0c\u5927\u5e45\u63d0\u5347\u5de5\u4f5c\u6548\u7387",
    ]:
        add_bullet(doc, h)

    # ===== 3. ARCH =====
    add_h(doc, "\u4e09\u3001\u7cfb\u7edf\u67b6\u6784", level=1)
    add_para(doc,
             "\u7cfb\u7edf\u91c7\u7528\u524d\u540e\u7aef\u5206\u79bb\u67b6\u6784\uff0c"
             "\u540e\u7aef\u57fa\u4e8e Python FastAPI \u6846\u67b6\u63d0\u4f9b RESTful API\uff0c"
             "\u6570\u636e\u5b58\u50a8\u4f7f\u7528 SQLite \u8f7b\u91cf\u6570\u636e\u5e93\uff1b"
             "\u524d\u7aef\u5206\u4e3a Web \u7ba1\u7406\u7aef\uff08\u5355\u9875\u5e94\u7528\uff09\u548c\u5fae\u4fe1\u539f\u751f\u5c0f\u7a0b\u5e8f\u7aef\u3002",
             size=11, sb=6, sa=8)

    make_table(doc, "\u7cfb\u7edf\u5c42\u6b21", [
        ("\u540e\u53f0\u7ba1\u7406\u7aef", "Web \u5355\u9875\u5e94\u7528\uff0c\u9879\u76ee\u7ba1\u7406\u3001Excel \u5bfc\u5165\u3001\u6587\u6863\u751f\u6210\u3001\u6570\u636e\u7edf\u8ba1"),
        ("\u5fae\u4fe1\u5c0f\u7a0b\u5e8f", "\u5fae\u4fe1\u539f\u751f\u5f00\u53d1\uff0c\u9879\u76ee\u6d4f\u89c8\u3001\u73b0\u573a\u62cd\u7167\u53d6\u8bc1\u3001\u6574\u6539\u4e0a\u4f20"),
        ("\u540e\u7aef\u670d\u52a1", "Python FastAPI\uff0c\u6570\u636e\u63a5\u53e3\u3001Excel \u89e3\u6790\u3001Word \u751f\u6210\u3001\u7167\u7247\u7ba1\u7406"),
        ("\u6570\u636e\u5b58\u50a8", "SQLite + \u6587\u4ef6\u7cfb\u7edf\uff0c\u9879\u76ee/\u9690\u60a3\u6570\u636e\u5e93\u3001\u7167\u7247\u4e0e\u6587\u6863\u6587\u4ef6\u5b58\u50a8"),
    ], accent="312E81")

    # ===== 4. ADMIN =====
    doc.add_page_break()
    add_h(doc, "\u56db\u3001\u540e\u53f0\u7ba1\u7406\u7aef\u529f\u80fd", level=1)
    add_para(doc,
             "\u540e\u53f0\u7ba1\u7406\u7aef\u662f\u4f9b\u7ba1\u7406\u4eba\u5458\u4f7f\u7528\u7684 Web \u7aef\u64cd\u4f5c\u754c\u9762\uff0c"
             "\u91c7\u7528\u73b0\u4ee3\u5316 UI \u8bbe\u8ba1\uff08\u7d2b\u9756\u8272\u4e3b\u9898\uff09\uff0c"
             "\u63d0\u4f9b\u6570\u636e\u603b\u89c8\u3001\u6e05\u5355\u5bfc\u5165\u3001\u9879\u76ee\u7ba1\u7406\u3001\u6587\u6863\u751f\u6210\u7b49\u6838\u5fc3\u529f\u80fd\u3002",
             size=11, sb=6, sa=6)
    add_para(doc, "\u8bbf\u95ee\u5730\u5740\uff1a" + ADMIN_URL, size=10, color=(99, 102, 241), sa=12)

    add_h(doc, "4.1  \u5de5\u4f5c\u53f0\uff08\u6570\u636e\u603b\u89c8\uff09", level=2)
    make_table(doc, "", [
        ("\u7edf\u8ba1\u5361\u7247", "\u5b9e\u65f6\u5c55\u793a\u9879\u76ee\u603b\u6570\u3001\u5df2\u5b8c\u6210\u3001\u8fdb\u884c\u4e2d\u3001\u5f85\u5904\u7406\u56db\u9879\u6838\u5fc3\u6307\u6807\uff0c\u4ee5\u53ca\u6240\u8986\u76d6\u7684\u8857\u9053\u6570\u91cf"),
        ("\u8857\u9053\u8fdb\u5ea6", "\u6309\u8857\u9053\u5206\u7ec4\u5c55\u793a\u68c0\u67e5\u5b8c\u6210\u8fdb\u5ea6\u6761\uff0c\u76f4\u89c2\u638c\u63e1\u5404\u8f96\u533a\u5de5\u4f5c\u8fdb\u5c55"),
        ("\u7167\u7247\u8fdb\u5ea6", "\u73af\u5f62\u56fe\u5c55\u793a\u5168\u5c40\u7167\u7247\u4e0a\u4f20\u767e\u5206\u6bd4\uff0c\u5feb\u901f\u4e86\u89e3\u6574\u4f53\u91c7\u96c6\u60c5\u51b5"),
        ("\u6700\u8fd1\u9879\u76ee", "\u8868\u683c\u5c55\u793a\u6700\u8fd1\u9879\u76ee\u7684\u540d\u79f0\u3001\u8857\u9053\u3001\u72b6\u6001\u3001\u9690\u60a3\u6570\u3001\u7167\u7247\u8fdb\u5ea6\uff0c\u652f\u6301\u5feb\u6377\u64cd\u4f5c"),
    ])

    add_h(doc, "4.2  \u4e0a\u4f20\u9690\u60a3\u6e05\u5355", level=2)
    make_table(doc, "", [
        ("\u6587\u4ef6\u4e0a\u4f20", "\u652f\u6301\u70b9\u51fb\u9009\u62e9\u6216\u62d6\u62fd\u4e0a\u4f20 Excel\uff08.xlsx\uff09\u683c\u5f0f\u7684\u9690\u60a3\u95ee\u9898\u6e05\u5355"),
        ("\u81ea\u52a8\u89e3\u6790", "\u7cfb\u7edf\u81ea\u52a8\u8bc6\u522b\u8857\u9053\u5206\u7ec4\u3001\u9879\u76ee\u4fe1\u606f\uff08\u540d\u79f0\u3001\u5404\u53c2\u5efa\u5355\u4f4d\u3001\u8054\u7cfb\u4eba\u7b49\uff09\u3001\u9690\u60a3\u660e\u7ec6"),
        ("\u7167\u7247\u63d0\u53d6", "\u81ea\u52a8\u4ece Excel \u5d4c\u5165\u56fe\u7247\u4e2d\u63d0\u53d6\u9690\u60a3\u7167\u7247\u5e76\u5173\u8054\u5230\u5bf9\u5e94\u8bb0\u5f55"),
        ("\u89e3\u6790\u62a5\u544a", "\u4e0a\u4f20\u5b8c\u6210\u540e\u5373\u65f6\u5c55\u793a\u89e3\u6790\u7ed3\u679c\uff1a\u9879\u76ee\u6570\u3001\u9690\u60a3\u6761\u6570\u3001\u7167\u7247\u6570\u7b49\u7edf\u8ba1"),
    ])

    add_h(doc, "4.3  \u9879\u76ee\u7ba1\u7406", level=2)
    make_table(doc, "", [
        ("\u7b5b\u9009\u67e5\u770b", "\u652f\u6301\u6309\u8857\u9053\u3001\u72b6\u6001\uff08\u5df2\u5b8c\u6210/\u8fdb\u884c\u4e2d/\u5f85\u5904\u7406\uff09\u7b5b\u9009\u9879\u76ee\u5217\u8868"),
        ("\u9879\u76ee\u5361\u7247", "\u6bcf\u4e2a\u9879\u76ee\u5361\u7247\u5c55\u793a\u540d\u79f0\u3001\u8857\u9053\u3001\u7c7b\u522b\u3001\u8054\u7cfb\u4eba\u3001\u9690\u60a3\u6570\u3001\u7167\u7247\u8fdb\u5ea6\u6761\u7b49\u4fe1\u606f"),
        ("\u6279\u91cf\u64cd\u4f5c", "\u5df2\u5b8c\u6210\u9879\u76ee\u53ef\u52fe\u9009\u540e\u6279\u91cf\u751f\u6210\u68c0\u67e5\u6863\u6848\u6587\u6863"),
        ("\u6587\u6863\u5217\u8868", "\u5c55\u793a\u6240\u6709\u5df2\u751f\u6210\u7684 Word \u6587\u6863\uff0c\u652f\u6301\u4e00\u952e\u4e0b\u8f7d"),
        ("\u7167\u7247\u91cd\u63d0", "\u5f53\u7167\u7247\u663e\u793a\u5f02\u5e38\u65f6\uff0c\u53ef\u4e00\u952e\u91cd\u65b0\u4ece\u539f\u59cb Excel \u63d0\u53d6\u9690\u60a3\u7167\u7247"),
    ])

    add_h(doc, "4.4  \u9879\u76ee\u8be6\u60c5", level=2)
    make_table(doc, "", [
        ("\u57fa\u7840\u4fe1\u606f", "\u5c55\u793a\u9879\u76ee\u540d\u79f0\u3001\u8857\u9053\u3001\u5730\u5740\u3001\u8d1f\u8d23\u4eba\u3001\u7535\u8bdd\u3001\u5efa\u8bbe/\u65bd\u5de5/\u76d1\u7406\u5355\u4f4d\u3001\u68c0\u67e5\u65e5\u671f"),
        ("\u9690\u60a3\u5217\u8868", "\u8868\u683c\u5c55\u793a\u5404\u6761\u9690\u60a3\u7684\u5e8f\u53f7\u3001\u7c7b\u578b\u3001\u63cf\u8ff0\u3001\u98ce\u9669\u7b49\u7ea7\u3001\u5206\u7c7b\u3001\u7167\u7247\u3001\u6574\u6539\u72b6\u6001"),
        ("\u73b0\u573a\u7167\u7247", "\u7f51\u683c\u5c55\u793a\u4f01\u4e1a\u95e8\u9762\u7167\u3001\u5ec9\u653f\u76d1\u7763\u5361\u7167\u3001\u6240\u6709\u9690\u60a3\u7167\u7247\u4e0e\u6574\u6539\u7167\u7247"),
        ("\u7167\u7247\u67e5\u770b", "\u70b9\u51fb\u7167\u7247\u53ef\u653e\u5927\u67e5\u770b\u539f\u56fe"),
    ])

    add_h(doc, "4.5  \u6587\u6863\u751f\u6210\u4e0e\u4e0b\u8f7d", level=2)
    make_table(doc, "", [
        ("\u4e00\u952e\u751f\u6210", "\u5bf9\u5df2\u5b8c\u6210\u9879\u76ee\u70b9\u51fb" + q("\u751f\u6210\u6587\u6863") + "\uff0c\u7cfb\u7edf\u81ea\u52a8\u586b\u5145 Word \u6a21\u677f\u5e76\u63d2\u5165\u6240\u6709\u7167\u7247"),
        ("\u6a21\u677f\u5957\u6253", "\u57fa\u4e8e\u9884\u8bbe Word \u6a21\u677f\uff0c\u81ea\u52a8\u586b\u5199\u5c01\u9762\u4fe1\u606f\u3001\u63d2\u5165\u95e8\u9762\u7167\u548c\u76d1\u7763\u5361\u3001\u751f\u6210\u9690\u60a3\u6574\u6539\u8bb0\u5f55\u8868"),
        ("\u6279\u91cf\u751f\u6210", "\u52fe\u9009\u591a\u4e2a\u5df2\u5b8c\u6210\u9879\u76ee\uff0c\u4e00\u952e\u5e76\u884c\u751f\u6210\u5168\u90e8\u6587\u6863"),
        ("\u6587\u6863\u4e0b\u8f7d", "\u751f\u6210\u5b8c\u6210\u540e\u53ef\u76f4\u63a5\u4e0b\u8f7d\u6807\u51c6 Word\uff08.docx\uff09\u683c\u5f0f\u7684\u68c0\u67e5\u6863\u6848"),
    ])

    # ===== 5. MINIAPP =====
    doc.add_page_break()
    add_h(doc, "\u4e94\u3001\u5fae\u4fe1\u5c0f\u7a0b\u5e8f\u7aef\u529f\u80fd", level=1)
    add_para(doc,
             "\u5fae\u4fe1\u5c0f\u7a0b\u5e8f\u7aef\u662f\u4f9b\u4e00\u7ebf\u68c0\u67e5\u4eba\u5458\u4f7f\u7528\u7684\u79fb\u52a8\u5de5\u5177\uff0c"
             "\u65e0\u9700\u5b89\u88c5 APP\uff0c\u5fae\u4fe1\u626b\u7801\u5373\u53ef\u4f7f\u7528\uff0c"
             "\u652f\u6301\u9879\u76ee\u6d4f\u89c8\u3001\u73b0\u573a\u62cd\u7167\u53d6\u8bc1\u548c\u6574\u6539\u7167\u7247\u4e0a\u4f20\u3002",
             size=11, sb=6, sa=12)

    add_h(doc, "5.1  \u68c0\u67e5\u4efb\u52a1\u5217\u8868", level=2)
    make_table(doc, "", [
        ("\u9879\u76ee\u5217\u8868", "\u6309\u8857\u9053\u5206\u7ec4\u5c55\u793a\u6240\u6709\u68c0\u67e5\u9879\u76ee\uff0c\u6bcf\u7ec4\u663e\u793a\u9879\u76ee\u6570\u91cf"),
        ("\u72b6\u6001\u6807\u7b7e", "\u6bcf\u4e2a\u9879\u76ee\u5c55\u793a" + q("\u5df2\u5b8c\u6210") + "/" + q("\u8fdb\u884c\u4e2d") + "/" + q("\u5f85\u5904\u7406") + "\u72b6\u6001\u6807\u7b7e\uff0c\u914d\u8272\u76f4\u89c2"),
        ("\u8fdb\u5ea6\u5c55\u793a", "\u6bcf\u4e2a\u9879\u76ee\u5c55\u793a\u7167\u7247\u4e0a\u4f20\u8fdb\u5ea6\u6761\u53ca\u6570\u5b57\uff08\u5982 3/5 \u56fe\u7247\uff09"),
        ("\u641c\u7d22\u529f\u80fd", "\u652f\u6301\u6309\u9879\u76ee\u540d\u79f0\u3001\u5730\u5740\u3001\u8054\u7cfb\u4eba\u5173\u952e\u5b57\u5b9e\u65f6\u641c\u7d22"),
        ("\u4e0b\u62c9\u5237\u65b0", "\u4e0b\u62c9\u9875\u9762\u5373\u53ef\u5237\u65b0\u6570\u636e\uff0c\u83b7\u53d6\u6700\u65b0\u72b6\u6001"),
    ], accent="059669")

    add_h(doc, "5.2  \u9879\u76ee\u8be6\u60c5", level=2)
    make_table(doc, "", [
        ("\u9879\u76ee\u4fe1\u606f", "\u5c55\u793a\u9879\u76ee\u540d\u79f0\u3001\u8857\u9053\u3001\u5730\u5740\u3001\u8054\u7cfb\u4eba\u3001\u7535\u8bdd\u3001\u7c7b\u522b\u3001\u5404\u53c2\u5efa\u5355\u4f4d\u7b49\u5b8c\u6574\u4fe1\u606f"),
        ("\u73b0\u573a\u7167\u7247", "\u5c55\u793a\u4f01\u4e1a\u95e8\u9762\u548c\u5ec9\u653f\u76d1\u7763\u5361\u4e24\u5f20\u73b0\u573a\u7167\u7247\uff0c\u70b9\u51fb\u53ef\u8df3\u8f6c\u62cd\u7167"),
        ("\u9690\u60a3\u5217\u8868", "\u5c55\u793a\u6240\u6709\u9690\u60a3\u8bb0\u5f55\uff0c\u542b\u5e8f\u53f7\u3001\u7c7b\u578b\u3001\u63cf\u8ff0\u6458\u8981\u3001\u9690\u60a3/\u6574\u6539\u7f29\u7565\u56fe\u3001\u72b6\u6001\u6807\u7b7e"),
        ("\u5feb\u6377\u64cd\u4f5c", "\u70b9\u51fb\u9690\u60a3\u6761\u76ee\u8fdb\u5165\u8be6\u60c5\u9875\uff1b\u70b9\u51fb\u7167\u7247\u533a\u57df\u8df3\u8f6c\u62cd\u7167\u9875\u9762"),
        ("\u4efb\u52a1\u63d0\u793a", "\u5e95\u90e8\u63d0\u793a\u5f53\u524d\u4efb\u52a1\u72b6\u6001\uff1a\u9700\u4e0a\u4f20\u7167\u7247\u6216\u4efb\u52a1\u5df2\u5b8c\u6210"),
    ], accent="059669")

    add_h(doc, "5.3  \u9690\u60a3\u8be6\u60c5\u4e0e\u6574\u6539", level=2)
    make_table(doc, "", [
        ("\u9690\u60a3\u4fe1\u606f", "\u5c55\u793a\u9690\u60a3\u7684\u5b8c\u6574\u63cf\u8ff0\u3001\u98ce\u9669\u7b49\u7ea7\uff08\u9ad8\u4eae\u6807\u6ce8\uff09\u3001\u5206\u7c7b\u3001\u53c2\u8003\u89c4\u8303\u3001\u5907\u6ce8"),
        ("\u9690\u60a3\u7167\u7247", "\u5c55\u793a\u4ece Excel \u5bfc\u5165\u7684\u539f\u59cb\u9690\u60a3\u7167\u7247\uff0c\u70b9\u51fb\u53ef\u9884\u89c8\u5927\u56fe"),
        ("\u6574\u6539\u4e0a\u4f20", "\u70b9\u51fb" + q("\u62cd\u7167\u4e0a\u4f20") + "\u6309\u94ae\u8df3\u8f6c\u76f8\u673a\u9875\u9762\uff0c\u62cd\u6444\u6574\u6539\u7167\u7247\u540e\u81ea\u52a8\u4e0a\u4f20"),
        ("\u91cd\u65b0\u62cd\u7167", "\u5df2\u4e0a\u4f20\u6574\u6539\u7167\u540e\u4ecd\u53ef\u70b9\u51fb" + q("\u91cd\u65b0\u62cd\u7167") + "\u8986\u76d6\u66f4\u65b0"),
        ("\u56fe\u7247\u9884\u89c8", "\u652f\u6301\u53cc\u56fe\uff08\u9690\u60a3\u7167+\u6574\u6539\u7167\uff09\u8054\u5408\u9884\u89c8\uff0c\u5de6\u53f3\u6ed1\u52a8\u5207\u6362"),
    ], accent="059669")

    add_h(doc, "5.4  \u73b0\u573a\u62cd\u7167\u53d6\u8bc1", level=2)
    make_table(doc, "", [
        ("\u53d6\u666f\u6846", "5:4 \u6a2a\u5411\u6784\u56fe\u53d6\u666f\u6846\uff0c\u4e0a\u4e0b\u9ed1\u8272\u906e\u7f69\u5f15\u5bfc\u62cd\u6444\u89d2\u5ea6"),
        ("\u9ad8\u6e05\u62cd\u6444", "\u8c03\u7528\u540e\u7f6e\u6444\u50cf\u5934\uff0c\u9ad8\u5206\u8fa8\u7387\u62cd\u6444\uff0c\u81ea\u52a8\u95ea\u5149\u706f"),
        ("\u667a\u80fd\u88c1\u5207", "\u62cd\u6444\u540e\u81ea\u52a8\u6309 5:4 \u6bd4\u4f8b\u88c1\u5207\uff0c\u786e\u4fdd\u6784\u56fe\u7edf\u4e00"),
        ("\u6c34\u5370\u53e0\u52a0", "\u4f01\u4e1a\u95e8\u9762\u548c\u76d1\u7763\u5361\u7167\u7247\u81ea\u52a8\u53e0\u52a0\u6c34\u5370\uff0c\u5305\u542b\u673a\u6784\u6807\u8bc6\u3001\u62cd\u6444\u65f6\u95f4\u3001\u9879\u76ee\u5730\u70b9"),
        ("\u9884\u89c8\u786e\u8ba4", "\u62cd\u6444\u540e\u9884\u89c8\u6548\u679c\uff0c\u786e\u8ba4\u65e0\u8bef\u540e\u518d\u4e0a\u4f20\uff0c\u652f\u6301\u91cd\u62cd"),
        ("\u81ea\u52a8\u4e0a\u4f20", "\u786e\u8ba4\u540e\u81ea\u52a8\u4e0a\u4f20\u81f3\u670d\u52a1\u5668\uff0c\u5168\u5c4f\u52a0\u8f7d\u63d0\u793a"),
    ], accent="059669")

    # ===== 6. TECH =====
    doc.add_page_break()
    add_h(doc, "\u516d\u3001\u6838\u5fc3\u6280\u672f\u4eae\u70b9", level=1)

    tech = [
        ("\u667a\u80fd Excel \u89e3\u6790",
         "\u91c7\u7528\u53cc\u901a\u9053\u56fe\u7247\u63d0\u53d6\u6280\u672f\uff08openpyxl \u5e93 + OOXML \u539f\u751f XML \u89e3\u6790\uff09\uff0c"
         "\u517c\u5bb9\u591a\u79cd Excel \u7248\u672c\u7684\u5d4c\u5165\u56fe\u7247\u683c\u5f0f\u3002\u652f\u6301\u5206\u6bb5\u5f0f\u653f\u52a1\u8868\u683c\u89e3\u6790\uff0c"
         "\u81ea\u52a8\u8bc6\u522b\u8857\u9053\u6807\u9898\u3001\u9879\u76ee\u4fe1\u606f\u5408\u5e76\u5355\u5143\u683c\u3001\u9690\u60a3\u6570\u636e\u884c\uff0c\u65e0\u9700\u4eba\u5de5\u6574\u7406\u3002"),
        ("Word \u6587\u6863\u81ea\u52a8\u751f\u6210",
         "\u57fa\u4e8e Word \u6a21\u677f\u52a8\u6001\u751f\u6210" + q("\u4e00\u4f01\u4e00\u6863") + "\u68c0\u67e5\u6863\u6848\u3002\u81ea\u52a8\u586b\u5145\u5c01\u9762\u4fe1\u606f\u3001"
         "\u63d2\u5165\u73b0\u573a\u7167\u7247\u3001\u6309\u9690\u60a3\u6570\u91cf\u52a8\u6001\u589e\u51cf\u8868\u683c\u884c\u5e76\u5d4c\u5165\u9690\u60a3\u7167/\u6574\u6539\u7167\uff0c"
         "\u7167\u7247\u6309\u6bd4\u4f8b\u81ea\u9002\u5e94\u7f29\u653e\uff0c\u8f93\u51fa\u5373\u4e3a\u6807\u51c6\u68c0\u67e5\u6863\u6848\u683c\u5f0f\u3002"),
        ("\u73b0\u573a\u53d6\u8bc1\u6c34\u5370",
         "\u5c0f\u7a0b\u5e8f\u7aef\u4f7f\u7528\u79bb\u5c4f Canvas \u6280\u672f\u5b9e\u65f6\u5408\u6210\u6c34\u5370\uff0c\u5305\u542b\u673a\u6784\u54c1\u724c\u6807\u8bc6\uff08SZEM\uff09\u3001"
         "\u673a\u6784\u540d\u79f0\u3001\u7cbe\u786e\u5230\u5206\u949f\u7684\u62cd\u6444\u65f6\u95f4\u3001\u9879\u76ee\u5730\u5740\uff0c\u786e\u4fdd\u53d6\u8bc1\u7167\u7247\u7684\u89c4\u8303\u6027\u4e0e\u53ef\u8ffd\u6eaf\u6027\u3002"),
        ("\u79fb\u52a8\u7aef\u56fe\u7247\u4f18\u5316",
         "\u91c7\u7528\u5206\u6279\u4e0b\u8f7d\u7b56\u7565\uff08\u6bcf\u6279 3 \u5f20\uff09\uff0c\u907f\u514d\u5e76\u53d1\u8fc7\u9ad8\u5bfc\u81f4\u52a0\u8f7d\u5931\u8d25\u3002"
         "\u6240\u6709\u8fdc\u7a0b\u56fe\u7247\u5148\u4e0b\u8f7d\u5230\u672c\u5730\u4e34\u65f6\u8def\u5f84\u518d\u5c55\u793a\uff0c\u517c\u5bb9\u5fae\u4fe1\u5c0f\u7a0b\u5e8f\u771f\u673a\u73af\u5883\u9650\u5236\u3002"),
        ("\u9879\u76ee\u72b6\u6001\u81ea\u52a8\u6d41\u8f6c",
         "\u7cfb\u7edf\u6839\u636e\u7167\u7247\u4e0a\u4f20\u60c5\u51b5\u81ea\u52a8\u66f4\u65b0\u9879\u76ee\u72b6\u6001\uff1a"
         "\u65e0\u4efb\u4f55\u4e0a\u4f20\u4e3a" + q("\u5f85\u5904\u7406") + "\uff0c\u90e8\u5206\u4e0a\u4f20\u4e3a" + q("\u8fdb\u884c\u4e2d") + "\uff0c"
         "\u6240\u6709\u6574\u6539\u7167 + \u95e8\u9762\u7167 + \u76d1\u7763\u5361\u9f50\u5168\u81ea\u52a8\u6807\u8bb0\u4e3a" + q("\u5df2\u5b8c\u6210") + "\u3002"),
        ("\u8f7b\u91cf\u5316\u90e8\u7f72",
         "\u540e\u7aef\u91c7\u7528 FastAPI + SQLite\uff0c\u65e0\u9700\u590d\u6742\u7684\u6570\u636e\u5e93\u8fd0\u7ef4\u3002"
         "\u7ba1\u7406\u7aef\u4e3a\u7eaf\u9759\u6001\u5355\u9875\u5e94\u7528\uff0c\u4e0e\u540e\u7aef\u540c\u6e90\u90e8\u7f72\uff0c\u5f00\u7bb1\u5373\u7528\u3002"),
    ]
    for title, desc in tech:
        add_h(doc, "\u25aa " + title, level=3, color=(79, 70, 229))
        add_para(doc, desc, size=10.5, sa=10)

    # ===== FOOTER =====
    doc.add_paragraph()
    add_para(doc, "\u2014 \u6587\u6863\u7ed3\u675f \u2014",
             size=10, color=(148, 163, 184),
             align=WD_ALIGN_PARAGRAPH.CENTER, sb=30)
    add_para(doc, "\u5982\u6709\u7591\u95ee\uff0c\u8bf7\u8054\u7cfb\u6280\u672f\u652f\u6301\u56e2\u961f",
             size=9, color=(148, 163, 184),
             align=WD_ALIGN_PARAGRAPH.CENTER)

    doc.save(OUTPUT)
    print("Done: " + OUTPUT)


if __name__ == "__main__":
    build()
